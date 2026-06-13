from __future__ import annotations

import argparse
import csv
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


def read_csv(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8", newline="") as f:
        return list(csv.DictReader(f))


def find_first_existing(paths: list[Path]) -> Path | None:
    for path in paths:
        if path.exists():
            return path
    return None


def pct(value: str | float) -> str:
    try:
        return f"{float(value) * 100:.1f}%"
    except Exception:
        return str(value)


def fmt_float(value: str | float, digits: int = 3) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def delta_pp(before: str | float, after: str | float) -> str:
    try:
        return f"{(float(after) - float(before)) * 100:+.2f} pp"
    except Exception:
        return ""


def normalize(text: str) -> str:
    return " ".join((text or "").lower().split())


def disagreement_score(row: dict) -> int:
    ref = normalize(row.get("text", ""))
    hyp = normalize(row.get("prediction", ""))
    if ref == hyp:
        return 0
    ref_words = set(ref.split())
    hyp_words = set(hyp.split())
    return len(ref_words.symmetric_difference(hyp_words)) + abs(len(ref) - len(hyp))


def count_prediction_changes(before_path: Path | None, after_path: Path | None) -> tuple[int, int]:
    if not before_path or not after_path or not before_path.exists() or not after_path.exists():
        return (0, 0)
    before = read_csv(before_path)
    after = read_csv(after_path)
    total = min(len(before), len(after))
    changed = sum(
        1 for i in range(total)
        if before[i].get("prediction", "") != after[i].get("prediction", "")
    )
    return changed, total


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        r.bold = True
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def setup_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    title = styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(11, 37, 69)

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def main() -> None:
    p = argparse.ArgumentParser(description="Build a midterm DOCX report from real pipeline outputs.")
    p.add_argument("--outputs_dir", default="outputs")
    p.add_argument("--out", default="outputs/midterm_report.docx")
    args = p.parse_args()

    outputs = Path(args.outputs_dir)
    dataset_stats = read_csv(outputs / "dataset_stats.csv")
    base_clean_path = outputs / "metrics_whisper_clean.csv"
    base_noisy_path = outputs / "metrics_whisper_noisy_by_snr.csv"
    forced_clean_path = find_first_existing([outputs / "metrics_whisper_clean_forced.csv", base_clean_path])
    forced_noisy_path = find_first_existing([outputs / "metrics_whisper_noisy_forced_by_snr.csv", base_noisy_path])
    base_clean_metrics = read_csv(base_clean_path)
    base_noisy_metrics = read_csv(base_noisy_path)
    clean_metrics = read_csv(forced_clean_path)
    noisy_metrics = read_csv(forced_noisy_path)
    noisy_predictions_path = find_first_existing([outputs / "whisper_noisy_forced.csv", outputs / "whisper_noisy.csv"])
    noisy_predictions = read_csv(noisy_predictions_path)
    clean_changed, clean_total = count_prediction_changes(outputs / "whisper_clean.csv", outputs / "whisper_clean_forced.csv")
    noisy_changed, noisy_total = count_prediction_changes(outputs / "whisper_noisy.csv", outputs / "whisper_noisy_forced.csv")

    clean_all = next(row for row in clean_metrics if row["group"] == "all")
    noisy_all = next(row for row in noisy_metrics if row["group"] == "all")
    base_clean_all = next(row for row in base_clean_metrics if row["group"] == "all")
    base_noisy_all = next(row for row in base_noisy_metrics if row["group"] == "all")
    noisy_by_snr = [row for row in noisy_metrics if row["group"] != "all"]
    noisy_by_snr.sort(key=lambda row: float(row["group"]), reverse=True)

    examples = [
        row for row in noisy_predictions
        if normalize(row.get("text", "")) != normalize(row.get("prediction", ""))
    ]
    examples.sort(key=disagreement_score, reverse=True)
    examples = examples[:3]

    doc = Document()
    setup_styles(doc)
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Midterm Progress Report").bold = True
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Tone-Aware LoRA Adaptation of PhoWhisper for Noise-Robust Vietnamese ASR").italic = True
    team = doc.add_paragraph()
    team.alignment = WD_ALIGN_PARAGRAPH.CENTER
    team.add_run("Mentor: Vu Ha Anh | Members: Nguyen Xuan Trung, Trinh Vy Kiet, Nguyen Thanh Phat, Pham Hoang Phuc")

    doc.add_heading("1. Executive Summary", level=1)
    doc.add_paragraph(
        "The midterm goal is to demonstrate real implementation progress rather than claim final model improvement. "
        "The current work has completed a reproducible Vietnamese noisy-ASR pipeline using VIVOS clean speech, MUSAN noise, "
        "controlled SNR mixing, Whisper-base zero-shot inference, forced Vietnamese decoding, and initial Vietnamese-specific metric prototypes."
    )
    add_bullet(doc, f"Clean baseline on 30 VIVOS utterances: WER {pct(clean_all['wer'])}, CER {pct(clean_all['cer'])}.")
    add_bullet(doc, f"Noisy baseline on 120 utterances across SNR 20/10/5/0 dB: WER {pct(noisy_all['wer'])}, CER {pct(noisy_all['cer'])}.")
    add_bullet(
        doc,
        "A source-code fix now passes language=vi and task=transcribe into Whisper generation; "
        f"it changed {clean_changed}/{clean_total} clean predictions and {noisy_changed}/{noisy_total} noisy predictions."
    )
    add_bullet(doc, "LoRA and decoder-side tone-aware multi-task learning are implemented as the next Colab training step, not yet claimed as an improvement.")

    doc.add_heading("2. Research Direction", level=1)
    doc.add_paragraph(
        "The project remains aligned with the proposal: evaluate whether explicit tone supervision can improve Vietnamese ASR robustness under controlled acoustic noise. "
        "The midterm scope is intentionally narrower: build the benchmark and baseline first, then compare Noisy LoRA against Tone-aware MTL after midterm."
    )
    add_bullet(doc, "Research question: Does explicit tone supervision improve Vietnamese ASR robustness beyond ordinary noisy LoRA fine-tuning?")
    add_bullet(doc, "Current safe claim: pipeline, baseline, and metric prototypes are complete enough for reproducible midterm evidence.")
    add_bullet(doc, "Not claimed yet: trained MTL superiority, lambda sensitivity, or enhancement-vs-tone preservation results.")

    doc.add_heading("3. Dataset and Benchmark", level=1)
    doc.add_paragraph(
        "VIVOS is used as the clean Vietnamese ASR source. MUSAN is used as the acoustic noise source. "
        "A fixed seed produces a controlled noisy test subset at SNR 20, 10, 5, and 0 dB."
    )
    add_table(
        doc,
        ["Manifest", "SNR", "Utterances", "Hours", "Avg. seconds"],
        [[row["manifest"], row["snr"], row["utterances"], row["hours"], row["avg_seconds"]] for row in dataset_stats],
    )

    doc.add_heading("4. Current Implementation", level=1)
    add_bullet(doc, "Manifest generation: VIVOS JSONL manifests and MUSAN noise manifest.")
    add_bullet(doc, "Noise benchmark: deterministic clean-noisy pairs with metadata for source audio, noise type, SNR, and seed.")
    add_bullet(doc, "Inference: Whisper-base zero-shot prediction CSV for clean and noisy subsets, with language=vi and task=transcribe forced at generation time.")
    add_bullet(doc, "Metrics: WER, CER, simple Tone Error Rate, and simple Diacritic Error Rate.")
    add_bullet(doc, "Prepared next step: PhoWhisper LoRA + decoder-side tone head with tone-label ignore-mask handling.")

    doc.add_heading("5. Preliminary Results", level=1)
    doc.add_paragraph(
        "The table below reports the current zero-shot Whisper-base baseline after forcing Vietnamese decoding. "
        "Values are error rates; lower is better. The before/after comparison shows that decoder control is useful but not enough by itself."
    )
    add_table(
        doc,
        ["Condition", "N", "WER before", "WER forced", "Delta", "CER before", "CER forced", "Delta"],
        [
            [
                "Clean",
                clean_all["n"],
                pct(base_clean_all["wer"]),
                pct(clean_all["wer"]),
                delta_pp(base_clean_all["wer"], clean_all["wer"]),
                pct(base_clean_all["cer"]),
                pct(clean_all["cer"]),
                delta_pp(base_clean_all["cer"], clean_all["cer"]),
            ],
            [
                "Noisy all",
                noisy_all["n"],
                pct(base_noisy_all["wer"]),
                pct(noisy_all["wer"]),
                delta_pp(base_noisy_all["wer"], noisy_all["wer"]),
                pct(base_noisy_all["cer"]),
                pct(noisy_all["cer"]),
                delta_pp(base_noisy_all["cer"], noisy_all["cer"]),
            ],
        ],
    )
    add_table(
        doc,
        ["Condition", "N", "WER", "CER", "TER simple", "DER simple"],
        [
            ["Clean", clean_all["n"], pct(clean_all["wer"]), pct(clean_all["cer"]), pct(clean_all["ter_simple"]), pct(clean_all["der_simple"])],
            ["Noisy all", noisy_all["n"], pct(noisy_all["wer"]), pct(noisy_all["cer"]), pct(noisy_all["ter_simple"]), pct(noisy_all["der_simple"])],
        ],
    )
    add_table(
        doc,
        ["SNR", "N", "WER", "CER", "TER simple", "DER simple"],
        [[row["group"], row["n"], pct(row["wer"]), pct(row["cer"]), pct(row["ter_simple"]), pct(row["der_simple"])] for row in noisy_by_snr],
    )
    doc.add_paragraph(
        "The strongest degradation appears at 0 dB, where WER increases to "
        f"{pct(next(row for row in noisy_by_snr if row['group'] == '0.0')['wer'])}. "
        "This supports the need for noise-robust adaptation and Vietnamese-specific analysis."
    )

    doc.add_heading("6. Qualitative Error Examples", level=1)
    for i, row in enumerate(examples, 1):
        doc.add_heading(f"Example {i}: SNR {row.get('snr')} dB, noise type {row.get('noise_type')}", level=2)
        doc.add_paragraph(f"Reference: {row.get('text', '')}")
        doc.add_paragraph(f"Prediction: {row.get('prediction', '')}")

    doc.add_heading("7. Risk Assessment", level=1)
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["CPU/local hardware", "Full training may be slow locally.", "Use Colab GPU for LoRA/MTL; keep local runs for preprocessing and inference subsets."],
            ["Tone-label noise", "Foreign words, acronyms, and numbers can corrupt tone supervision.", "Use ignore-mask policy and validate coverage on sample transcripts."],
            ["Overclaiming", "Midterm may be challenged if MTL has not trained.", "State clearly that current evidence is pipeline + baseline; MTL comparison is next."],
            ["Metric maturity", "TER/DER are prototypes and not fully edit-aligned yet.", "Report them as prototype metrics and improve alignment after midterm."],
        ],
    )

    doc.add_heading("8. Next Steps After Midterm", level=1)
    add_bullet(doc, "Run PhoWhisper zero-shot on the same clean/noisy subset.")
    add_bullet(doc, "Train clean LoRA and noisy LoRA on Colab using batch size 1, FP16 where available, and gradient accumulation.")
    add_bullet(doc, "Train tone-aware MTL and compare against noisy LoRA with WER, CER, TER, and DER.")
    add_bullet(doc, "Replace simple TER/DER with edit-aligned syllable-level metrics.")
    add_bullet(doc, "Optionally evaluate enhancement only after the core LoRA-vs-MTL comparison is stable.")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.add_run("Artifacts used: ").bold = True
    note.add_run(
        "outputs/dataset_stats.csv, outputs/whisper_clean.csv, outputs/whisper_noisy.csv, "
        "outputs/whisper_clean_forced.csv, outputs/whisper_noisy_forced.csv, "
        "outputs/metrics_whisper_clean.csv, outputs/metrics_whisper_noisy_by_snr.csv, "
        "outputs/metrics_whisper_clean_forced.csv, outputs/metrics_whisper_noisy_forced_by_snr.csv."
    )

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print(f"wrote {out}")


if __name__ == "__main__":
    main()
